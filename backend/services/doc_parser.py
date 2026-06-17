import io
import pdfplumber


SUPPORTED_EXTENSIONS = (".pdf", ".md", ".txt", ".docx", ".xlsx", ".xls")


def _extract_docx(content: bytes) -> str:
    """Pull text from a .docx: paragraphs + table cells, in document order."""
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _extract_xlsx(content: bytes) -> str:
    """Pull text from an .xlsx/.xls: each sheet, rows as tab-separated values."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append("\t".join(cells))
    wb.close()
    return "\n".join(parts)


def has_meaningful_text(text: str, min_chars: int = 50) -> bool:
    """
    True if the extracted text looks like real document content rather than
    near-empty output from a scanned/image-only PDF (which pdfplumber can't
    OCR). Feeding such near-empty text to the LLM produces hallucinated or
    empty extractions with no indication of why.
    """
    return len(text.strip()) >= min_chars


def extract_text_from_file(filename: str, content: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    elif lower.endswith((".md", ".txt")):
        return content.decode("utf-8", errors="replace")
    elif lower.endswith(".docx"):
        return _extract_docx(content)
    elif lower.endswith((".xlsx", ".xls")):
        return _extract_xlsx(content)
    else:
        raise ValueError(f"Unsupported file type: {filename}")
